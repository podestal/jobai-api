"""
Utility functions for text extraction from resume files.
"""
import os
import logging
from io import BytesIO

logger = logging.getLogger(__name__)


def extract_text_from_file(file):
    """
    Extract text from uploaded resume file.
    Supports PDF, DOC, and DOCX formats.
    
    Args:
        file: Django UploadedFile object
        
    Returns:
        str: Extracted text, or empty string if extraction fails
    """
    file_name = file.name.lower()
    file.seek(0)  # Reset file pointer to beginning
    
    try:
        if file_name.endswith('.pdf'):
            return extract_text_from_pdf(file)
        elif file_name.endswith('.docx'):
            return extract_text_from_docx(file)
        elif file_name.endswith('.doc'):
            return extract_text_from_doc(file)
        else:
            logger.warning(f"Unsupported file type: {file_name}")
            return ""
    except Exception as e:
        logger.error(f"Error extracting text from {file_name}: {str(e)}")
        # Return empty string instead of raising exception to not break the upload
        return ""


def extract_text_from_pdf(file):
    """Extract text from PDF file using pdfplumber (preferred) or PyPDF2 as fallback"""
    file.seek(0)
    
    try:
        # Try pdfplumber first (better text extraction)
        import pdfplumber
        with pdfplumber.open(BytesIO(file.read())) as pdf:
            text_parts = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n\n".join(text_parts)
    except ImportError:
        logger.warning("pdfplumber not available, falling back to PyPDF2")
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {str(e)}, trying PyPDF2")
    
    # Fallback to PyPDF2
    try:
        import PyPDF2
        file.seek(0)
        pdf_reader = PyPDF2.PdfReader(BytesIO(file.read()))
        text_parts = []
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"PyPDF2 extraction failed: {str(e)}")
        return ""


def extract_text_from_docx(file):
    """Extract text from DOCX file"""
    try:
        from docx import Document
        file.seek(0)
        doc = Document(BytesIO(file.read()))
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        return ""


def extract_text_from_doc(file):
    """
    Extract text from DOC file (older Microsoft Word format).
    Note: DOC files are binary and harder to parse.
    This is a basic implementation - you might want to use textract or antiword for better results.
    """
    try:
        # Try using python-docx (might work for some DOC files)
        from docx import Document
        file.seek(0)
        doc = Document(BytesIO(file.read()))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.warning(f"DOC extraction failed (DOC format is complex): {str(e)}")
        # DOC files are binary format and require special libraries like textract or antiword
        # For now, return empty string
        return ""

